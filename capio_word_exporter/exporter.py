"""Module for generating docx file"""
from docx import Document
from docx.shared import RGBColor


def _format_time(seconds):
    """Convert time """
    sec = seconds % 60
    seconds = seconds // 60
    minute = int(seconds % 60)
    seconds = seconds // 60
    hour = int(seconds % 60)
    return '{hour:02d}:{minute:02d}:{sec:05.2f}'.format(
        hour=hour, minute=minute, sec=sec)


def _add_start_time(paragraph, start_time):
    run = paragraph.add_run(_format_time(start_time))
    run.bold = True
    run.font.color.rgb = RGBColor(101, 98, 149)


def _add_word(paragraph, word, color=None, **_):
    run = paragraph.add_run(word)
    if color:
        run.font.color.rgb = RGBColor(*color)


def _add_sentence(doc, datum):
    para = doc.add_paragraph()
    _add_start_time(para, datum['time']['start'])
    para.add_run('\t')
    for word in datum['words']:
        para.add_run(' ')
        _add_word(para, **word)


def gen_docx(data, outputpath):
    """Generate docx file from parsed transcript data

    Parameters
    ----------
    data : list
        List of sentence data. See `parser.parse_transcript` for the detail.

    outputpath : str
        Path to save docx data.
    """
    doc = Document()
    for datum in data:
        _add_sentence(doc, datum)
    doc.save(outputpath)
